import random
import json
import os
from const import DATA_DIR


HEADING = '## '

QUESTIONS = {
    'sspace': 'Should “shared spaces” in urban planning be promoted?',
    'biohacking': 'Should greater regulatory control be exerted over genetic biohacking?'
}

# initially generated with random.shuffle
ORDERINGS = [
    {'graph': [10, 23, 33, 31, 32, 28, 17, 19, 26, 20, 22, 11, 21, 27, 15, 24, 34, 12, 25, 30, 16, 18, 29, 14],
     'plain': [33, 19, 21, 14, 16, 29, 12, 23, 31, 15, 28, 34, 20, 17, 24, 32, 27, 11, 30, 26, 25, 10, 18, 22]},
    {'graph': [26, 16, 31, 22, 10, 34, 15, 18, 20, 33, 29, 14, 12, 24, 27, 28, 21, 32, 23, 19, 30, 25, 17, 11],
     'plain': [29, 21, 30, 33, 34, 18, 11, 26, 24, 12, 19, 25, 31, 20, 17, 28, 10, 32, 27, 15, 22, 14, 23, 16]},
    {'graph': [34, 10, 22, 32, 28, 29, 33, 23, 30, 26, 11, 31, 25, 14, 19, 20, 12, 24, 16, 17, 15, 21, 18, 27],
     'plain': [15, 29, 27, 25, 10, 17, 20, 30, 24, 23, 16, 19, 18, 33, 34, 14, 21, 32, 28, 31, 22, 11, 12, 26]}
]

CONDITION_ORDERINGS = [
    ["graph", "plain"],
    ["plain", "graph"],
    ["plain", "graph"],
]


def get_full_ordering(marker_num):
    full_ordering = []
    condition_ordering = CONDITION_ORDERINGS[marker_num]
    num_per_condition = len(ORDERINGS[marker_num]['graph'])
    for i in range(num_per_condition):
        for condition in [0, 1]:
            full_ordering.append(ORDERINGS[marker_num][condition_ordering[condition]][i])

    all_condition_ordering = condition_ordering * num_per_condition

    return full_ordering, all_condition_ordering


def get_question(cond, graph_is_first, sspace_is_first):
    is_first = (cond == "graph" and graph_is_first) or (cond == "plain" and not graph_is_first)
    if is_first:
        if sspace_is_first:
            return "sspace"
        else:
            return "biohacking"
    else:
        if sspace_is_first:
            return "biohacking"
        else:
            return "sspace"

        
def parse_arg(arg):
    return arg.replace('&#x27;', "'").replace('&quot;', '"').replace('&gt;', '>')
        

def get_args():
    args = {'graph': {}, 'plain': {}}

    for filename in os.listdir(DATA_DIR):
        with open(os.path.join(DATA_DIR, filename), 'r') as f:
            data = json.load(f)
            arg = parse_arg(data['argument'])
            params = data['params']
            question = get_question(params['condition'], params['novelToolFirst'], params['sspaceFirst'])

            if int(params['experimentId']) in args[params['condition']].keys():
                print(f"ERROR: duplicate at {params['condition']}, num {params['experimentId']}")

            args[params['condition']][int(params['experimentId'])] = {'arg': arg, 'question': question}

    print(f'files: {len(os.listdir(DATA_DIR))}')

    for c, cond in args.items():
        sspaces = [t for t in cond.values() if t['question'] == 'sspace']
        biohacking = [t for t in cond.values() if t['question'] == 'biohacking']
        print(f'tot {len(cond.values())} {c}: sspaces: {len(sspaces)}, bio: {len(biohacking)}')

    return args


def get_orderings(args):
    orderings = {}
    for cond in args:
        ordering = list(args[cond].keys())
        random.shuffle(ordering)
        orderings[cond] = ordering

    return orderings


def add_arg(doc, arg, question):
    doc.add_paragraph(QUESTIONS[question], style="Heading 2")

    for para in arg.split('\n'):
        is_heading = para.startswith(HEADING)
        if is_heading:
            para = para[len(HEADING):]

        if para:
            style = "Heading 3" if is_heading else "Normal"
            doc.add_paragraph(para, style=style)

    table = doc.add_table(rows=2, cols=5, style="Table Grid")
    header_texts = ['Clarity', 'Persuasiveness', 'Structure', 'Objection responsiveness', 'Overall assessment']
    headers = table.rows[0].cells
    assert(len(headers) == len(header_texts))
    for i, header in enumerate(headers):
        header.text = header_texts[i]

    for cell in table.rows[1].cells:
        cell.text = "       /10"


def generate_doc(args, orderings, first_cond, second_cond):
    document = docx.Document()

    for i in range(len(orderings['graph'])):
        for cond in [first_cond, second_cond]:
            add_arg(document, **args[cond][orderings[cond][i]])

            if cond == first_cond or i < len(orderings['graph']) - 1:
                document.add_page_break()

    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.3)

    return document


def generate_all_docs():
    args = get_args()
    assert (len(args['graph']) == len(args['plain']))

    for i in range(3):
        document_generated = generate_doc(args, ORDERINGS[i], *CONDITION_ORDERINGS[i])

        doc_filename = f'marker_sheet_{i+1}.docx'
        try:
            os.remove(doc_filename)
        except FileNotFoundError:
            pass

        document_generated.save(doc_filename)


if __name__ == '__main__':
    import docx
    from docx.shared import Inches
    
    generate_all_docs()
